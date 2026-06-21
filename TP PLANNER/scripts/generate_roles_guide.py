# -*- coding: utf-8 -*-
"""Génère le document Word : rôle de chaque page — TP Planner."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Roles_des_pages_TP_Planner.docx"

PAGES = [
    {
        "title": "Page d'accueil publique",
        "fichier": "index.php → pages/landing.php",
        "acces": "Tout visiteur (non connecté)",
        "role": (
            "Vitrine du site. Elle présente l'application TP Planner, ses avantages "
            "et les liens vers la connexion ou l'inscription. Elle ne modifie aucune "
            "donnée : son rôle est uniquement informatif et d'orientation des utilisateurs."
        ),
    },
    {
        "title": "Connexion",
        "fichier": "login.php / pages/login.php",
        "acces": "Tout visiteur",
        "role": (
            "Point d'entrée sécurisé de l'application. Elle authentifie l'utilisateur "
            "(e-mail + mot de passe), crée la session PHP et redirige vers l'espace adapté : "
            "tableau de bord administrateur ou espace laboratoire du stagiaire."
        ),
    },
    {
        "title": "Inscription stagiaire",
        "fichier": "register.php / pages/register.php",
        "acces": "Tout visiteur",
        "role": (
            "Permet aux professeurs stagiaires de créer leur propre compte en choisissant "
            "une classe existante. Elle enregistre un nouvel enregistrement dans la table "
            "students ; l'administrateur n'a pas besoin de créer le compte manuellement."
        ),
    },
    {
        "title": "Tableau de bord administrateur",
        "fichier": "pages/dashboard.php",
        "acces": "Administrateur",
        "role": (
            "Centre de pilotage pour l'admin. Affiche en un coup d'œil les statistiques "
            "(classes, séances TP, quiz), les éléments récents, les alertes sur les "
            "checklists incomplètes et un graphique des scores de quiz par classe."
        ),
    },
    {
        "title": "Gestion des classes",
        "fichier": "pages/classes.php",
        "acces": "Administrateur",
        "role": (
            "Référentiel des groupes / niveaux scolaires (ex. « 2 BAC SM 1 »). "
            "L'admin y crée, modifie et supprime les classes. Chaque classe sert ensuite "
            "à rattacher les stagiaires et les séances TP."
        ),
    },
    {
        "title": "Liste des séances TP",
        "fichier": "pages/tp_sessions.php",
        "acces": "Administrateur",
        "role": (
            "Inventaire de toutes les fiches de travaux pratiques. Permet de rechercher, "
            "filtrer par classe, trier et accéder rapidement à la création, la modification, "
            "la consultation ou la suppression d'une séance."
        ),
    },
    {
        "title": "Édition d'une séance TP",
        "fichier": "pages/tp_edit.php",
        "acces": "Administrateur",
        "role": (
            "Atelier de rédaction de la fiche TP. C'est ici que l'admin construit le contenu "
            "pédagogique : informations générales, étapes, matériel, checklist par phase, "
            "questions du mini-quiz, schéma et import de document. C'est la page de production."
        ),
    },
    {
        "title": "Consultation d'une séance TP",
        "fichier": "pages/tp_view.php",
        "acces": "Administrateur et stagiaire (sa classe uniquement)",
        "role": (
            "Page de lecture et d'utilisation en séance. Affiche la fiche complète, "
            "permet de cocher la checklist en temps réel et de passer le mini-quiz. "
            "L'admin y voit aussi les scores de tous les stagiaires ; le stagiaire ne voit que les siens."
        ),
    },
    {
        "title": "Export PDF",
        "fichier": "pages/tp_pdf.php",
        "acces": "Utilisateur connecté ayant accès à la séance",
        "role": (
            "Génère un document PDF imprimable à partir d'une fiche TP (titre, objectifs, "
            "étapes, matériel, checklist, quiz). Sert à l'archivage, l'impression ou le partage "
            "hors ligne ; ne modifie pas la base de données."
        ),
    },
    {
        "title": "Gestion des comptes",
        "fichier": "pages/stagiaires.php",
        "acces": "Administrateur",
        "role": (
            "Administration des utilisateurs du système : création, modification, suppression "
            "des comptes administrateur (users) et stagiaire (students), attribution d'une "
            "classe et changement de mot de passe ou de rôle."
        ),
    },
    {
        "title": "Espace laboratoire (stagiaire)",
        "fichier": "pages/member_dashboard.php",
        "acces": "Professeur stagiaire",
        "role": (
            "Tableau de bord du stagiaire. Affiche sa classe et la liste des TP qui lui "
            "sont destinés. C'est le point de départ pour ouvrir une fiche ou consulter "
            "ses résultats de quiz."
        ),
    },
    {
        "title": "Résultats de quiz (stagiaire)",
        "fichier": "pages/quiz_results.php",
        "acces": "Professeur stagiaire",
        "role": (
            "Bilan personnel des réponses au mini-quiz. Pour une séance choisie, affiche "
            "chaque question, la réponse donnée, la bonne réponse et le score — pour permettre "
            "la révision et l'auto-évaluation après la séance."
        ),
    },
    {
        "title": "Changement de langue",
        "fichier": "pages/set_lang.php",
        "acces": "Appelé par le sélecteur FR / EN / AR",
        "role": (
            "Page technique sans interface visible. Enregistre la langue choisie (français, "
            "anglais ou arabe) en session et cookie, puis renvoie l'utilisateur sur la page "
            "d'où il venait avec les textes traduits."
        ),
    },
    {
        "title": "Déconnexion",
        "fichier": "pages/logout.php",
        "acces": "Utilisateur connecté",
        "role": (
            "Ferme proprement la session utilisateur et redirige vers la page d'accueil "
            "publique. Garantit qu'aucune donnée personnelle ne reste accessible après "
            "quittage de l'application."
        ),
    },
]


def para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    return p


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

    h = doc.add_heading("TP Planner", 0)
    for r in h.runs:
        r.font.name = "Calibri"
    para(doc, "Rôle de chaque page du site", bold=True, size=14)
    para(
        doc,
        "Ce document décrit le rôle fonctionnel de chaque écran : "
        "à quoi sert la page, qui peut y accéder et ce qu'elle fait dans l'application.",
    )
    para(doc, "Application : gestion des travaux pratiques (TP) en laboratoire")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("Sommaire", 1)
    for i, p in enumerate(PAGES, 1):
        para(doc, f"{i}. {p['title']} — {p['fichier']}")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    for i, page in enumerate(PAGES, 1):
        doc.add_heading(f"{i}. {page['title']}", 1)
        para(doc, "Fichier : " + page["fichier"], bold=True)
        para(doc, "Qui y accède : " + page["acces"])
        doc.add_heading("Rôle de la page", 2)
        para(doc, page["role"])
        if i < len(PAGES):
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Créé : {OUT}")


if __name__ == "__main__":
    build()
