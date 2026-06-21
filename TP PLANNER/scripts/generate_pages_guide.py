# -*- coding: utf-8 -*-
"""Génère le guide Word des pages du site TP Planner."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Guide_des_pages_TP_Planner.docx"

PAGES = [
    {
        "title": "1. Page d'accueil publique",
        "url": "index.php (→ pages/landing.php si non connecté)",
        "acces": "Public — aucune connexion requise",
        "resume": (
            "Première page du site lorsqu'on n'est pas connecté. Présente TP Planner "
            "avec un design laboratoire (animations, visuels). Sections : accueil (hero), "
            "à propos, fonctionnalités clés, contact avec formulaire. Liens vers connexion "
            "et inscription. Sélecteur de langue FR / EN / AR. Les coordonnées peuvent "
            "être lues depuis la table site_settings."
        ),
        "captures": [
            "Vue complète du hero (titre, sous-titre, boutons Connexion / Inscription)",
            "Section À propos et cartes fonctionnalités",
            "Section Contact et pied de page",
        ],
    },
    {
        "title": "2. Connexion",
        "url": "login.php ou pages/login.php",
        "acces": "Public",
        "resume": (
            "Formulaire de connexion par e-mail et mot de passe. Vérifie les comptes "
            "dans users (administrateurs) puis students (professeurs stagiaires). "
            "Redirige vers le tableau de bord admin ou l'espace laboratoire selon le rôle. "
            "Messages d'erreur traduits (champs vides, compte inconnu, etc.). Lien vers "
            "l'inscription stagiaire."
        ),
        "captures": [
            "Formulaire vide",
            "Message d'erreur (optionnel)",
        ],
    },
    {
        "title": "3. Inscription (professeur stagiaire)",
        "url": "register.php ou pages/register.php",
        "acces": "Public",
        "resume": (
            "Permet à un nouveau professeur stagiaire de créer un compte : nom, e-mail, "
            "mot de passe (confirmation), choix de la classe parmi celles créées par "
            "l'administrateur. Après succès, redirection vers la page de connexion. "
            "Bloque l'inscription si aucune classe n'existe encore."
        ),
        "captures": [
            "Formulaire d'inscription avec liste déroulante des classes",
            "Message de succès après inscription (sur la page login)",
        ],
    },
    {
        "title": "4. Tableau de bord administrateur",
        "url": "pages/dashboard.php",
        "acces": "Administrateur connecté",
        "resume": (
            "Page d'accueil après connexion admin. Affiche des cartes statistiques : "
            "nombre de classes, séances TP, questions de quiz. Alerte si des checklists "
            "restent incomplètes. Tableaux des classes et séances récentes. Graphique "
            "Chart.js des scores moyens de quiz par classe. Accès rapide vers la gestion "
            "des stagiaires."
        ),
        "captures": [
            "Cartes statistiques et alerte checklist",
            "Graphique des scores et listes récentes",
        ],
    },
    {
        "title": "5. Gestion des classes",
        "url": "pages/classes.php",
        "acces": "Administrateur",
        "resume": (
            "Liste, création, modification et suppression des classes (groupes / niveaux). "
            "Chaque classe peut être associée à un enseignant référent (admin). Recherche "
            "par nom. Lien pour voir les séances TP filtrées par classe."
        ),
        "captures": [
            "Liste des classes avec actions",
            "Formulaire d'ajout ou de modification",
        ],
    },
    {
        "title": "6. Liste des séances TP",
        "url": "pages/tp_sessions.php",
        "acces": "Administrateur",
        "resume": (
            "Vue d'ensemble de toutes les fiches de travaux pratiques. Filtres : recherche "
            "texte, classe, tri par titre / date / durée. Actions : créer, modifier, voir, "
            "supprimer une séance. La suppression efface aussi matériel, checklist, quiz "
            "et réponses associées."
        ),
        "captures": [
            "Tableau avec filtres et tri",
            "Boutons Voir / Modifier / Supprimer",
        ],
    },
    {
        "title": "7. Création et édition d'une séance TP",
        "url": "pages/tp_edit.php (avec ou sans ?id=)",
        "acces": "Administrateur",
        "resume": (
            "Formulaire complet pour une fiche TP : titre, numéro de fiche, unité, "
            "objectifs et compétences (éditeur Quill), consignes de sécurité, durée, "
            "classe, schéma image, import document PDF/Word. Onglets ou sections pour "
            "ajouter les étapes, le matériel (par type), les items de checklist "
            "(avant / pendant / après) et les questions du mini-quiz avec bonne réponse. "
            "Enregistrement par étapes avec protection CSRF."
        ),
        "captures": [
            "En-tête de la fiche (infos générales)",
            "Section étapes / matériel / checklist / quiz",
            "Import de document ou schéma",
        ],
    },
    {
        "title": "8. Consultation d'une séance TP",
        "url": "pages/tp_view.php?id=…",
        "acces": "Administrateur ou stagiaire de la classe concernée",
        "resume": (
            "Affichage lecture de la fiche TP : objectifs, compétences, sécurité, "
            "étapes, matériel, document importé, schéma. Checklist interactive : "
            "cocher/décocher les tâches par phase. Mini-quiz : le stagiaire répond "
            "aux QCM (son nom est pris automatiquement) ; l'admin peut saisir un nom "
            "et voir tous les scores. Bouton vers l'export PDF pour les utilisateurs "
            "autorisés."
        ),
        "captures": [
            "Corps de la fiche (objectifs, étapes)",
            "Checklist avec cases à cocher",
            "Bloc quiz et tableau des scores",
        ],
    },
    {
        "title": "9. Export PDF d'une séance",
        "url": "pages/tp_pdf.php?id=…",
        "acces": "Utilisateur connecté avec accès à la séance",
        "resume": (
            "Génère et télécharge un fichier PDF de la fiche TP (TCPDF) : en-tête "
            "avec titre, classe, durée, contenu structuré (objectifs, étapes, matériel, "
            "checklist, questions du quiz). Nécessite composer install. S'ouvre "
            "directement dans le navigateur ou se télécharge selon le réglage du navigateur."
        ),
        "captures": [
            "Aperçu du PDF généré (première page)",
            "Page suivante avec checklist ou quiz",
        ],
    },
    {
        "title": "10. Gestion des comptes (stagiaires et admins)",
        "url": "pages/stagiaires.php",
        "acces": "Administrateur",
        "resume": (
            "Administration des utilisateurs : liste des professeurs stagiaires (table "
            "students) et des administrateurs (table users). Ajout, modification "
            "(nom, e-mail, classe pour stagiaire, rôle, mot de passe), suppression. "
            "Permet de promouvoir un stagiaire en admin ou l'inverse."
        ),
        "captures": [
            "Liste des comptes avec badges de rôle",
            "Formulaire d'ajout ou d'édition",
        ],
    },
    {
        "title": "11. Espace laboratoire (tableau de bord stagiaire)",
        "url": "pages/member_dashboard.php",
        "acces": "Professeur stagiaire connecté",
        "resume": (
            "Accueil du professeur stagiaire après connexion. Affiche le groupe / classe "
            "rattaché au compte et la liste des séances TP publiées pour ce groupe. "
            "Liens pour ouvrir une fiche TP ou consulter les résultats de quiz pour "
            "une séance donnée. Message d'avertissement si le profil ou la classe manque."
        ),
        "captures": [
            "Carte « Mon groupe »",
            "Tableau des séances avec boutons Ouvrir / Résultats quiz",
        ],
    },
    {
        "title": "12. Mes résultats de quiz",
        "url": "pages/quiz_results.php (?tp_id= optionnel)",
        "acces": "Professeur stagiaire",
        "resume": (
            "Historique détaillé des réponses du stagiaire connecté. Sélection d'une "
            "séance TP dans une liste déroulante. Pour chaque question : énoncé, réponse "
            "choisie, bonne réponse, indication correct/incorrect et score total. "
            "Utile pour réviser après avoir passé le mini-quiz sur tp_view."
        ),
        "captures": [
            "Sélecteur de séance TP",
            "Tableau détail question par question avec score global",
        ],
    },
    {
        "title": "13. Changement de langue",
        "url": "pages/set_lang.php",
        "acces": "Tous (via sélecteur dans l'interface)",
        "resume": (
            "Point technique appelé par le sélecteur FR / EN / AR. Enregistre la langue "
            "en session et cookie puis redirige vers la page précédente. Pas d'interface "
            "visuelle propre : documenter plutôt le sélecteur dans la barre de navigation."
        ),
        "captures": [
            "Sélecteur de langue dans la navbar (admin ou stagiaire)",
            "Même sélecteur sur la page d'accueil",
        ],
    },
    {
        "title": "14. Déconnexion",
        "url": "pages/logout.php",
        "acces": "Utilisateur connecté",
        "resume": (
            "Détruit la session et les cookies de session, puis redirige vers la page "
            "d'accueil publique. Accessible depuis le menu utilisateur (icône profil) "
            "→ Déconnexion."
        ),
        "captures": [
            "Menu déroulant avec lien Déconnexion",
            "Page d'accueil après déconnexion",
        ],
    },
]


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Page de garde
    t = doc.add_heading("TP Planner", 0)
    for run in t.runs:
        run.font.name = "Calibri"
    add_para(doc, "Guide des pages du site web", bold=True)
    add_para(
        doc,
        "Ce document décrit chaque écran de l'application en français. "
        "Utilisez les encadrés « Captures suggérées » pour placer vos screenshots.",
    )
    add_para(doc, "Projet : TP PLANNER — Gestion des travaux pratiques en laboratoire")
    add_para(doc, "Date de génération : mai 2026")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # Table des matières manuelle
    add_heading(doc, "Table des matières", 1)
    for page in PAGES:
        add_para(doc, page["title"])
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # Chaque page
    for i, page in enumerate(PAGES):
        add_heading(doc, page["title"], 1)
        add_para(doc, "URL / fichier : ", bold=True)
        add_para(doc, page["url"])
        add_para(doc, "Accès : ", bold=True)
        add_para(doc, page["acces"])
        add_heading(doc, "Résumé", 2)
        add_para(doc, page["resume"])
        add_heading(doc, "Captures d'écran suggérées", 2)
        for cap in page["captures"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(cap)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        add_para(doc, "[ Insérer votre capture d'écran ici ]")
        if i < len(PAGES) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build()
